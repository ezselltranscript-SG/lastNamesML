#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Servicio para procesar documentos Word y corregir apellidos mal escritos.
Versión mejorada con detección avanzada de apellidos y corrección de dos etapas.
"""

import os
import sys
import argparse
import logging
import re
import time
from datetime import datetime
from typing import List, Dict, Tuple, Any, Optional

import pandas as pd
import numpy as np
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

# Configurar logging
log_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'word_service_debug.log')
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file, mode='w'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger('word_service')

# Add the project root to the path so we can import our modules
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# Configure better error reporting
import traceback

from src.models import load_model, SurnameCorrector
from src.data_processing import load_reference_data
from src.surname_detection import extract_potential_surnames, load_exclusion_lists
from src.feature_engineering import extract_features


def parse_arguments():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(description='Correct surnames in Word documents using advanced detection')
    parser.add_argument('--model-path', required=True, help='Path to the trained correction model')
    parser.add_argument('--surname-model-path', help='Path to the trained surname detection model')
    parser.add_argument('--reference-data', required=True, help='Path to reference data with correct surnames')
    parser.add_argument('--input-file', help='Path to input Word document')
    parser.add_argument('--input-dir', help='Path to directory with input Word documents')
    parser.add_argument('--output-dir', default='data/corrected_word', help='Directory to save corrected documents')
    parser.add_argument('--confidence-threshold', type=float, default=0.7, 
                        help='Confidence threshold for corrections (0-1)')
    parser.add_argument('--surname-confidence-threshold', type=float, default=0.6,
                        help='Confidence threshold for surname detection (0-1)')
    parser.add_argument('--use-ner', action='store_true', default=True,
                        help='Use Named Entity Recognition for surname detection')
    parser.add_argument('--places-exclusion-file', help='Path to file with place names to exclude')
    parser.add_argument('--words-exclusion-file', help='Path to file with common words to exclude')
    
    args = parser.parse_args()
    
    # Validate arguments
    if not args.input_file and not args.input_dir:
        parser.error("Either --input-file or --input-dir must be provided")
    
    # Validate input paths
    if args.input_file and not os.path.exists(args.input_file):
        parser.error(f"Input file does not exist: {args.input_file}")
    
    if args.input_dir and not os.path.exists(args.input_dir):
        parser.error(f"Input directory does not exist: {args.input_dir}")
    
    # Create output directory if it doesn't exist
    os.makedirs(args.output_dir, exist_ok=True)
    
    return args


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from a Word document.
    
    Args:
        docx_path: Path to Word document
        
    Returns:
        Extracted text
    """
    try:
        doc = Document(docx_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from {docx_path}: {e}")
        logger.error(traceback.format_exc())
        return ""


def process_word_document(docx_path: str, corrector: SurnameCorrector, output_dir: str, 
                       confidence_threshold: float = 0.7,
                       surname_confidence_threshold: float = 0.6) -> Tuple[str, str, List[Dict[str, Any]]]:
    """
    Process a Word document to correct surnames using advanced detection.
    
    Args:
        docx_path: Path to Word document
        corrector: Surname corrector instance
        output_dir: Directory to save corrected document
        confidence_threshold: Confidence threshold for corrections
        surname_confidence_threshold: Confidence threshold for surname detection
        
    Returns:
        Tuple of (corrected_doc_path, report_path, corrections)
    """
    try:
        logger.info(f"Procesando documento: {docx_path}")
        
        # Load document
        logger.debug("Cargando documento...")
        doc = Document(docx_path)
        
        # Extract text from document
        logger.debug("Extrayendo texto...")
        full_text = extract_text_from_docx(docx_path)
        logger.debug(f"Texto extraído: {len(full_text)} caracteres")
        
        # Detect and correct surnames
        logger.info("Detectando y corrigiendo apellidos...")
        _, corrections = corrector.correct_text(full_text, confidence_threshold, surname_confidence_threshold)
        logger.info(f"Correcciones encontradas: {len(corrections)}")
        
        # Log detailed information about each correction
        for i, correction in enumerate(corrections):
            logger.debug(f"Corrección {i+1}: {correction['original']} -> {correction['corrected']} "  
                        f"(Confianza detección: {correction['surname_confidence']:.2f}, "  
                        f"Confianza corrección: {correction['correction_confidence']:.2f})")
            logger.debug(f"Contexto: '{correction['context']}'")
        
        # Apply corrections to document
        logger.debug("Aplicando correcciones al documento...")
        corrected_doc = apply_corrections_to_document(doc, corrections)
        
        # Save corrected document
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        corrected_doc_path = os.path.join(output_dir, f"{base_name}_corregido.docx")
        logger.info(f"Guardando documento corregido: {corrected_doc_path}")
        corrected_doc.save(corrected_doc_path)
        
        # Generate report
        report_path = os.path.join(output_dir, f"{base_name}_informe.docx")
        logger.info(f"Generando informe: {report_path}")
        generate_report(docx_path, corrected_doc_path, corrections, report_path)
        
        return corrected_doc_path, report_path, corrections
        
    except Exception as e:
        logger.error(f"Error procesando documento {docx_path}: {e}")
        logger.error(traceback.format_exc())
        raise


def apply_corrections_to_document(doc: Document, corrections: List[Dict[str, Any]]) -> Document:
    """
    Apply corrections to a Word document.
    
    Args:
        doc: Word document
        corrections: List of corrections
        
    Returns:
        Corrected document
    """
    # Create a new document for the corrected version
    corrected_doc = Document()
    
    # Copy document properties
    for section in doc.sections:
        new_section = corrected_doc.add_section()
        new_section.page_height = section.page_height
        new_section.page_width = section.page_width
        new_section.left_margin = section.left_margin
        new_section.right_margin = section.right_margin
        new_section.top_margin = section.top_margin
        new_section.bottom_margin = section.bottom_margin
    
    # Create a mapping of corrections by position
    correction_map = {}
    for correction in corrections:
        if 'position' in correction:
            correction_map[correction['position']] = correction
    
    # Process each paragraph
    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            # Add empty paragraph to corrected document
            corrected_doc.add_paragraph()
            continue
        
        # Create a new paragraph for the corrected text
        new_para = corrected_doc.add_paragraph()
        new_para.style = para.style
        
        # Apply corrections if any match this paragraph
        current_pos = 0
        for correction in corrections:
            original = correction['original']
            corrected = correction['corrected']
            
            # Find all occurrences of the original text in this paragraph
            start_pos = text.find(original, current_pos)
            if start_pos == -1:
                continue
                
            # Add text before the correction
            if start_pos > current_pos:
                new_para.add_run(text[current_pos:start_pos])
            
            # Add the corrected text with highlighting
            run = new_para.add_run(corrected)
            run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
            
            # Update current position
            current_pos = start_pos + len(original)
        
        # Add any remaining text
        if current_pos < len(text):
            new_para.add_run(text[current_pos:])
        elif current_pos == 0:  # No corrections were applied
            new_para.text = text
    
    # Copy tables
    for table in doc.tables:
        new_table = corrected_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.style = table.style
        
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.cell(i, j).text = cell.text
    
    return corrected_doc


def generate_report(original_doc_path: str, corrected_doc_path: str, 
                 corrections: List[Dict[str, Any]], report_path: str) -> None:
    """
    Generate a detailed report of surname detections and corrections.
    
    Args:
        original_doc_path: Path to original document
        corrected_doc_path: Path to corrected document
        corrections: List of corrections
        report_path: Path to save report
    """
    try:
        # Create report document
        report_doc = Document()
        
        # Add title
        title = report_doc.add_heading('Informe de Corrección de Apellidos', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add document information
        report_doc.add_heading('Información del Documento', level=1)
        report_doc.add_paragraph(f'Documento original: {os.path.basename(original_doc_path)}')
        report_doc.add_paragraph(f'Documento corregido: {os.path.basename(corrected_doc_path)}')
        report_doc.add_paragraph(f'Fecha de procesamiento: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Add summary
        report_doc.add_heading('Resumen', level=1)
        report_doc.add_paragraph(f'Total de correcciones: {len(corrections)}')
        
        # Add corrections table
        if corrections:
            report_doc.add_heading('Correcciones Detalladas', level=1)
            table = report_doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Apellido Original'
            header_cells[1].text = 'Apellido Corregido'
            header_cells[2].text = 'Confianza Detección'
            header_cells[3].text = 'Confianza Corrección'
            header_cells[4].text = 'Contexto'
            
            # Add data rows
            for correction in corrections:
                row_cells = table.add_row().cells
                row_cells[0].text = correction['original']
                row_cells[1].text = correction['corrected']
                row_cells[2].text = f"{correction.get('surname_confidence', 'N/A'):.2f}"
                row_cells[3].text = f"{correction.get('correction_confidence', correction.get('confidence', 0.0)):.2f}"
                
                # Mostrar un extracto del contexto (limitado para que quepa en la tabla)
                context = correction.get('context', '')
                if len(context) > 50:
                    context = context[:47] + '...'
                row_cells[4].text = context
        else:
            report_doc.add_paragraph('No se encontraron apellidos para corregir.')
        
        # Añadir sección de metodología
        report_doc.add_heading('Metodología', level=1)
        methodology_text = (
            "El sistema utiliza un enfoque de dos etapas para la corrección de apellidos:\n\n"
            "1. Detección de Apellidos: Se utiliza una combinación de Named Entity Recognition (NER) "
            "y reglas contextuales para identificar palabras que probablemente sean apellidos.\n\n"
            "2. Corrección de Apellidos: Para cada apellido detectado, se utiliza un modelo de "
            "machine learning entrenado para determinar la versión correcta del apellido.\n\n"
            "La 'Confianza Detección' indica la seguridad del sistema de que la palabra es un apellido. "
            "La 'Confianza Corrección' indica la seguridad en la corrección propuesta."
        )
        report_doc.add_paragraph(methodology_text)
        
        # Save report
        report_doc.save(report_path)
        logger.info(f"Informe guardado: {report_path}")
        
    except Exception as e:
        logger.error(f"Error generando informe: {e}")
        logger.error(traceback.format_exc())
        raise


def process_word_directory(dir_path: str, corrector: SurnameCorrector, output_dir: str,
                        confidence_threshold: float = 0.7,
                        surname_confidence_threshold: float = 0.6) -> List[Tuple[str, str, List[Dict[str, Any]]]]:
    """
    Process all Word documents in a directory.
    
    Args:
        dir_path: Path to directory with Word documents
        corrector: Surname corrector instance
        output_dir: Directory to save corrected documents
        confidence_threshold: Confidence threshold for corrections
        surname_confidence_threshold: Confidence threshold for surname detection
        
    Returns:
        List of tuples (corrected_doc_path, report_path, corrections)
    """
    results = []
    
    try:
        # Get all Word documents in directory
        docx_files = [f for f in os.listdir(dir_path) 
                     if f.lower().endswith('.docx') and not f.lower().endswith('_corregido.docx')]
        
        logger.info(f"Encontrados {len(docx_files)} documentos Word")
        
        # Process each document
        for docx_file in docx_files:
            docx_path = os.path.join(dir_path, docx_file)
            try:
                result = process_word_document(
                    docx_path, 
                    corrector, 
                    output_dir, 
                    confidence_threshold,
                    surname_confidence_threshold
                )
                results.append(result)
                logger.info(f"Completado: {docx_file} - {len(result[2])} correcciones")
            except Exception as e:
                logger.error(f"Error procesando {docx_file}: {e}")
                logger.error(traceback.format_exc())
        
        logger.info("\nProcesamiento de directorio completado con éxito!")
        return results
        
    except Exception as e:
        logger.error(f"Error procesando directorio {dir_path}: {e}")
        logger.error(traceback.format_exc())
        return results


def main():
    """Main function to run the Word document correction service."""
    try:
        logger.info("\n" + "=" * 50)
        logger.info("Servicio de Corrección de Apellidos en Documentos Word")
        logger.info("=" * 50)
        
        args = parse_arguments()
        
        # Mostrar argumentos recibidos
        logger.info(f"Modelo: {args.model_path}")
        logger.info(f"Datos de referencia: {args.reference_data}")
        logger.info(f"Documento de entrada: {args.input_file}")
        logger.info(f"Directorio de salida: {args.output_dir}")
        
        # Verificar que los archivos existen
        logger.info("Verificando archivos de entrada:")
        if args.input_file:
            if not os.path.exists(args.input_file):
                logger.error(f"ERROR: El archivo de entrada no existe: {args.input_file}")
                return 1
            logger.info(f"  Archivo de entrada OK: {args.input_file}")
        
        if not os.path.exists(args.model_path):
            logger.error(f"ERROR: El modelo no existe: {args.model_path}")
            return 1
            
        if not os.path.exists(args.reference_data):
            logger.error(f"ERROR: Los datos de referencia no existen: {args.reference_data}")
            return 1
            
        # Load reference data
        logger.info("Cargando datos de referencia...")
        reference_df = load_reference_data(args.reference_data)
        logger.info(f"Datos de referencia cargados: {len(reference_df)} registros")
        
        # Extract surnames from reference data
        if 'CleanedSurname' in reference_df.columns:
            all_surnames = reference_df['CleanedSurname'].dropna().unique().tolist()
        elif 'ExtractedSurname' in reference_df.columns:
            all_surnames = reference_df['ExtractedSurname'].dropna().unique().tolist()
        elif 'surname' in reference_df.columns:
            all_surnames = reference_df['surname'].dropna().unique().tolist()
        elif 'full_name' in reference_df.columns:
            _, all_surnames = extract_names_surnames(reference_df, name_col='full_name')
        else:
            logger.error("ERROR: Los datos de referencia deben contener una columna con apellidos")
            return 1
        
        logger.info(f"Cargados {len(all_surnames)} apellidos de referencia")
        logger.info(f"Ejemplos de apellidos: {all_surnames[:5]}...")
        
        # Cargar listas de exclusión (lugares y palabras comunes)
        logger.info("Cargando listas de exclusión...")
        common_places, common_words = load_exclusion_lists()
        logger.info(f"Cargadas {len(common_places)} lugares y {len(common_words)} palabras comunes para exclusión")
        
        # Load the correction model
        logger.info("Cargando modelo de corrección...")
        correction_model = load_model(args.model_path)
        logger.info(f"Modelo de corrección cargado: {type(correction_model).__name__}")
        
        # Load the surname detection model if available
        surname_detection_model = None
        if args.surname_model_path and os.path.exists(args.surname_model_path):
            logger.info("Cargando modelo de detección de apellidos...")
            try:
                surname_detection_model = load_model(args.surname_model_path)
                logger.info(f"Modelo de detección de apellidos cargado: {type(surname_detection_model).__name__}")
            except Exception as e:
                logger.warning(f"No se pudo cargar el modelo de detección de apellidos: {e}")
                logger.warning("Se utilizará detección basada en reglas")
        else:
            logger.info("No se especificó modelo de detección de apellidos. Usando detección basada en reglas.")
        
        # Create surname corrector
        logger.info("Inicializando corrector de apellidos...")
        corrector = SurnameCorrector()
        corrector.correction_model = correction_model
        corrector.is_surname_model = surname_detection_model
        corrector.has_surname_classifier = surname_detection_model is not None
        corrector.set_reference_surnames(all_surnames)
        corrector.set_exclusion_lists(common_places, common_words)
        
        # Configurar uso de NER
        corrector.use_ner = args.use_ner
        
        # Process input file or directory
        if args.input_file:
            logger.info(f"Procesando archivo: {args.input_file}")
            process_word_document(
                args.input_file, 
                corrector, 
                args.output_dir, 
                args.confidence_threshold,
                args.surname_confidence_threshold
            )
        elif args.input_dir:
            logger.info(f"Procesando directorio: {args.input_dir}")
            process_word_directory(
                args.input_dir, 
                corrector, 
                args.output_dir, 
                args.confidence_threshold,
                args.surname_confidence_threshold
            )

            

        
        logger.info("\nProcesamiento completado con éxito!")
        return 0
        
    except Exception as e:
        logger.error("\nError durante la ejecución:")
        logger.error(str(e))
        logger.error("\nDetalles completos del error:")
        logger.error(traceback.format_exc())
        return 1


if __name__ == "__main__":
    # Import here to avoid circular imports
    from src.data_processing import extract_names_surnames
    
    main()
